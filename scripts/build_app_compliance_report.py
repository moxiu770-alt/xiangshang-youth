from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("文档/向上少年App上线手续与备案合规分析报告.docx")

# Preset: compact_reference_guide. Named override: China A4 portrait for a mainland
# compliance report; typography, spacing, tables, and furniture otherwise follow
# the preset's compact reference guide tokens.
FONT = "Microsoft YaHei"
FONT_LATIN = "Aptos"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
GREEN = "1F6B4F"
BLACK = "202124"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D0D5DD", size="6"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def set_run_font(run, name=FONT, size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(p, before=0, after=6, line=1.25, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        pPr = p._p.get_or_add_pPr()
        k = OxmlElement("w:keepNext")
        pPr.append(k)


def add_p(doc, text="", style=None, before=0, after=6, line=1.25, align=None, size=11, color=BLACK, bold=False, italic=False, keep=False):
    p = doc.add_paragraph(style=style)
    set_para(p, before, after, line, align, keep)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_p(doc, parts, before=0, after=6, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align, keep)
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run_font(r, **kwargs)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    before = {1: 18, 2: 14, 3: 10}[level]
    after = {1: 10, 2: 7, 3: 5}[level]
    p = doc.add_paragraph()
    set_para(p, before, after, 1.15, keep=True)
    r = p.add_run(text)
    set_run_font(r, size=sizes[level], color=colors[level], bold=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.left_indent = Inches(0.375 if level == 0 else 0.625)
    pf.first_line_indent = Inches(-0.188)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.left_indent = Inches(0.375)
    pf.first_line_indent = Inches(-0.188)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)
    return p


def add_note(doc, label, text, fill=CALLOUT, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA], indent=120)
    set_table_borders(table, color="D5DCE6", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.25)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    add_p(doc, "", after=3)
    return table


def add_table(doc, headers, rows, widths, font_size=9.5, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        set_para(p, 0, 0, 1.1, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i >= len(cells):
                break
            p = cells[i].paragraphs[0]
            set_para(p, 0, 0, 1.18, WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    add_p(doc, "", after=3)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def setup_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        s._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        s.font.size = Pt(11)
    # Quiet running header and footer, suitable for a serious internal report.
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = "向上少年 App 上线手续与备案合规分析报告"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(hp, 0, 0, 1.0)
    for r in hp.runs:
        set_run_font(r, size=9, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def build():
    doc = Document()
    setup_document(doc)
    core = doc.core_properties
    core.title = "向上少年 App 上线手续与备案合规分析报告"
    core.subject = "中国大陆 App 上线手续、备案及功能触发许可分析"
    core.author = ""
    core.keywords = "App备案;教育App;未成年人;个人信息保护;等保;医疗器械"

    # Cover / masthead
    add_p(doc, "项目合规报告", before=22, after=10, size=11, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "向上少年 App 上线手续与备案合规分析报告", before=0, after=6, size=25, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "中国大陆公开上架、校园部署及真实学生数据处理场景", before=0, after=18, size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    meta = [
        ("报告版本", "V1.0"),
        ("分析基准日", "2026年8月11日"),
        ("发布地区", "中国大陆"),
        ("运营主体", "待确认"),
        ("适用范围", "iOS、Android、学校管理后台及配套接口"),
    ]
    add_table(doc, ["项目字段", "内容"], meta, [2300, 7060], font_size=10.5, header_fill=LIGHT_GRAY)
    add_note(doc, "结论先行", "当前版本不能直接面向公众或学校上线。它至少同时触发 App 备案、教育 App 备案、网络安全等级保护、公安联网备案、未成年人个人信息保护和应用商店审核；若保留“AI诊断”、付费培训、第三方商城、公开班级圈或医疗服务，还会进一步触发专项许可或更高合规要求。", fill="EEF4FB", color=BLUE)
    add_p(doc, "本报告基于现有产品方案、当前 App 功能描述及中国大陆现行公开法规政策整理，供项目立项、开发整改、备案申报和上线评审使用。涉及医疗器械、地方体育培训、网络出版等事项，最终以运营地主管部门书面确认或正式审批结果为准。", before=12, after=8, size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_p(doc, "内部使用 · 合规评审稿", before=16, after=0, size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    doc.add_page_break()

    add_heading(doc, "一、产品合规画像", 1)
    add_p(doc, "现有方案显示，App 设有家庭端、教师端和校长端，服务对象包括学校、教师、家长及学生；功能覆盖学生绑定、体测任务、运动能力评估、视力/口腔/心理筛查入口、姿态与步态照片或视频、健康档案、报告、课程建议、班级动态、客服和预约等。该组合决定了上线审核不能只按“运动工具”处理。", after=8)
    add_table(doc, ["功能或数据", "合规属性", "上线影响"], [
        ("学校、班级、学生档案", "教育场景及未成年人个人信息", "教育 App 备案、最小权限、学校数据协议"),
        ("身高、体重、BMI、体测成绩", "健康信息，属于敏感个人信息", "单独同意、影响评估、严格访问控制"),
        ("站姿、前屈、步态照片/视频", "图像、身体影像；可能含生物特征", "相机权限、原始素材保存期限、删除机制"),
        ("AI评分、报告、风险标签", "自动化分析及可能对学生有重大影响的结果", "透明说明、人工复核、避免歧视性标签"),
        ("课程、预约、商城", "可能构成收费服务或交易平台", "消费者权益、支付规则、地方许可或 EDI"),
        ("班级圈、评论、客服消息", "用户发布内容和互动服务", "实名、审核、举报、日志、应急响应"),
    ], [2500, 3000, 3860], font_size=9.5)

    add_heading(doc, "二、必须在上线前完成的事项", 1)
    add_p(doc, "以下项目属于公开上架和真实校园部署的基础门槛。没有完成的项目，应列为上线阻断项。", after=8)
    add_table(doc, ["序号", "事项", "办理/建设要求", "责任与状态"], [
        ("1", "运营主体", "确定中国大陆公司作为 App 主办者、个人信息处理者、应用商店开发者和收款主体；营业执照范围与实际业务一致。", "甲方/运营主体\n待确认"),
        ("2", "域名、云和生产环境", "域名实名认证；生产 API、数据库、文件/视频存储采用合规云服务；固定生产域名、IP、SSL、包名和签名。", "技术/运营\n未完成"),
        ("3", "网站 ICP 备案", "对官网、隐私政策页和生产域名办理非经营性 ICP 备案；若构成经营性互联网信息服务，另行核定 ICP 许可证。", "运营/云服务商\n未完成"),
        ("4", "App 备案", "iOS、Android App 分别提交工信部 App 备案；包名、应用名、主体、域名和备案信息必须一致。", "运营/云服务商\n未完成"),
        ("5", "教育 App 提供者备案", "向教育主管部门系统提交运营主体、应用信息、服务场景和安全材料；每所学校按要求办理使用者备案。", "运营主体/学校\n未完成"),
        ("6", "等级保护", "完成系统定级、备案、整改和必要测评；后台、App API、学校管理端和数据系统应纳入范围。", "技术/安全服务商\n未完成"),
        ("7", "公安联网备案", "网络正式联通后30日内，通过全国互联网安全管理服务平台办理公安联网备案，并按要求维护变更信息。", "运营/技术\n未完成"),
        ("8", "隐私与儿童保护", "正式协议、儿童隐私规则、监护人同意、敏感信息单独同意、撤回、删除、注销、投诉和数据导出均须可用。", "产品/法务/技术\n未完成"),
        ("9", "学校数据协议", "与学校明确处理者/受托处理者关系、授权范围、保存期限、分权访问、事件通报及终止后返还或删除。", "运营/学校\n未完成"),
        ("10", "内容安全", "班级圈、照片/视频、客服、评论等建立实名、审核、举报、处置、留痕和应急响应闭环。", "产品/运营\n未完成"),
        ("11", "应用商店材料", "企业开发者账号、签名证书、隐私标签、权限说明、年龄分级、审核账号、备案号、截图和真实可用后台。", "发布负责人\n未完成"),
    ], [600, 1900, 5100, 1760], font_size=9.2)

    add_heading(doc, "三、未成年人个人信息与数据安全要求", 1)
    add_p(doc, "本项目处理学生姓名、学校、班级、健康、体测、照片/视频、心理或风险提示等数据。未满十四周岁信息和健康信息均应按高保护等级设计。学校提供数据时，不等于平台可以无限制使用。", after=8)
    add_table(doc, ["控制点", "上线要求"], [
        ("监护人授权", "建立可验证的父母或其他监护人同意流程；不得用一个概括勾选框替代儿童数据授权。"),
        ("单独同意", "健康信息、原始照片/视频、向第三方心理系统提供数据等分别说明并取得单独同意。"),
        ("数据最小化", "能在本地完成的姿态提取尽量本地完成；没有必要时不上传原始视频，不采集与功能无关的数据。"),
        ("访问控制", "学校、年级、班级、教师和家长按租户及角色隔离；校长不得默认查看不必要的学生原始素材。"),
        ("保存与删除", "建立原始视频、姿态关键点、报告、日志和备份的保存期限；注销、撤回或目的完成后删除或停止处理。"),
        ("第三方处理", "云服务、短信、推送、AI服务商、心理测评系统分别签订数据处理协议，禁止无授权转委托。"),
        ("影响评估", "对敏感信息、自动化决策、学校批量导入、第三方共享和跨境传输分别形成书面评估记录。"),
        ("年度审计", "每年自行或委托专业机构开展未成年人个人信息保护合规审计，并保留报告和整改证据。"),
    ], [2200, 7160], font_size=10)
    add_note(doc, "当前代码整改", "现有客户端有用户协议、隐私政策和儿童隐私政策入口，但展示内容仍为简短占位说明；同时尚未发现完整的账号注销、数据导出、删除和撤回同意闭环。该部分不能以“后续补充”方式带入正式上线。", fill="FFF8E7", color=GOLD)

    doc.add_page_break()
    add_heading(doc, "四、按功能触发的专项许可或备案", 1)
    add_table(doc, ["功能保留情况", "可能触发事项", "建议"], [
        ("保留“AI诊断”、疾病判断、康复或治疗建议", "医疗器械软件分类界定、注册/备案；如开展在线诊疗，还涉及医疗机构及互联网诊疗资质。", "首发改为“运动能力评估/体测辅助分析”，去除诊断和治疗表述。"),
        ("直接招收未成年人并销售体育培训", "所在地体育类非学科校外培训准入、登记、收费和预付费监管。", "首发仅做学校 SaaS 或与持证培训机构合作。"),
        ("商城允许第三方商家入驻", "EDI 在线数据处理与交易处理业务许可及平台经营者责任。", "首发采用自营商品或关闭商城。"),
        ("App 内销售数字课程", "iOS 内购规则、各 Android 商店数字商品政策、退款和自动续费规则。", "区分数字内容、实体商品和线下服务的支付路径。"),
        ("公开发布教材、音视频读物", "可能涉及网络出版服务许可或视听节目服务许可。", "核验版权；必要时与持证内容平台合作。"),
        ("班级圈扩展为公开社区、评论或直播", "网络信息内容服务、用户实名、内容审核、投诉举报；具有舆论或社会动员属性时还需安全评估。", "首发做学校封闭圈，关闭公开发布和直播。"),
        ("基于画像个性化推荐课程", "算法推荐透明、关闭个性化推荐、标签管理及未成年人保护；不当然等于算法备案。", "保留人工推荐和非个性化选项。"),
        ("使用境外云、境外分析 SDK 或境外 AI", "个人信息出境合规、告知同意、标准合同/认证/安全评估等条件。", "首发全部使用境内供应商并清理不必要 SDK。"),
    ], [2600, 3900, 2860], font_size=9.2)

    add_heading(doc, "五、当前版本上线阻断项", 1)
    add_p(doc, "以下问题不是一般优化项，而是会影响备案、商店审核或真实学生数据使用的 P0 问题。", after=8)
    blockers = [
        "后台仍以 Mock 数据为主，真实账户、验证码、权限、学校租户隔离、接口鉴权和生产数据生命周期尚未完成。",
        "法律文书入口虽已存在，但用户协议、隐私政策和儿童隐私政策仍是占位文本，缺少处理清单、第三方 SDK、保存期限和权利行使方式。",
        "没有完整的监护人可验证授权、健康信息单独同意、照片/视频单独同意、撤回授权、账号注销、数据导出和删除闭环。",
        "身体影像、步态视频、课堂照片及姿态关键点的保存期限、原始素材删除和备份清理机制尚未确定。",
        "心理测评第三方系统的运营主体、服务资质、数据接收范围、授权文案和结果回传协议尚未确定。",
        "班级圈、客服、课程上传等用户内容功能缺少完整内容审核后台、举报处置、日志留存和应急响应。",
        "“综合运动能力 AI 诊断平台”“专业身心测评与科学健康干预”“风险学生”等表述可能引发医疗、未成年人标签和自动化决策风险。",
        "生产主体、域名、签名、应用商店开发者账号、ICP/APP备案号、教育 App 备案和等保材料尚未形成。",
    ]
    for item in blockers:
        add_bullet(doc, item)
    add_note(doc, "测试数据要求", "在上述整改完成前，不应录入真实学生数据，包括学校内部试点。测试阶段使用虚拟数据、脱敏数据或经过充分匿名化的数据。", fill="FDECEC", color=RED)

    add_heading(doc, "六、推荐的分阶段上线方案", 1)
    add_heading(doc, "第一阶段：学校封闭试点版", 2)
    add_p(doc, "目标是先验证学校端、教师端、家长端和体测辅助流程，控制许可范围。建议只保留学校账户、学生绑定、体测任务、运动能力评估、报告展示和非医疗课程建议。", after=6)
    for item in [
        "关闭公开班级圈、公开评论、直播、公开视频上传和第三方商家商城。",
        "关闭付费数字课程和面向个人家长的培训招生；课程先作为学校服务内容。",
        "关闭医疗诊断、疾病筛查、治疗、康复和在线医生咨询表述。",
        "心理模块只做明确授权后的外部跳转，首期不自动回传敏感结果，或改为由合规第三方独立处理。",
        "全部使用境内云和境内 SDK，采用虚拟或经授权的最小化数据。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "第二阶段：商业化功能扩展", 2)
    add_p(doc, "在第一阶段完成备案、学校协议、隐私审计和安全运行后，再逐项开启收费课程、专家预约、商城、公开内容和个性化推荐。每增加一类功能，都应重新进行功能定性、隐私影响评估和应用商店资质复核。", after=6)

    add_heading(doc, "七、标准办理流程", 1)
    steps = [
        ("1", "确定主体与产品定位", "确定运营公司、App名称、商标、收费模式、数据控制者；先决定是否走非医疗路线。"),
        ("2", "固定生产资源", "完成域名实名认证、国内云、API、数据库、文件存储、包名、Bundle ID、签名和生产环境。"),
        ("3", "完成合规设计", "完成数据地图、权限清单、SDK清单、隐私政策、儿童规则、监护人授权、学校数据协议和删除策略。"),
        ("4", "完成安全建设", "完成等保定级、备案、整改及必要测评；完成渗透测试、接口鉴权、租户隔离、日志和应急预案。"),
        ("5", "办理基础备案", "办理网站 ICP 备案、教育 App 提供者备案、App 备案，并准备各学校使用者备案。具体顺序与省级主管部门确认。"),
        ("6", "办理专项资质", "按最终功能办理体育培训、ICP许可证、EDI、医疗器械、网络出版或视听相关资质。"),
        ("7", "发布前验证", "完成隐私检测、权限测试、儿童授权测试、账号注销测试、支付/退款测试、异常和删除测试。"),
        ("8", "应用商店审核", "提交 Apple App Store 及各 Android 应用市场；备案主体、应用名、隐私政策、截图和实际功能必须一致。"),
        ("9", "上线后维护", "网络正式联通后30日内办理公安联网备案；持续维护备案变更、未成年人年度审计、等保和内容安全记录。"),
    ]
    add_table(doc, ["阶段", "步骤", "主要工作"], steps, [700, 2400, 6260], font_size=9.5)

    doc.add_page_break()
    add_heading(doc, "八、申报材料清单", 1)
    add_table(doc, ["类别", "应准备材料"], [
        ("主体材料", "营业执照、统一社会信用代码、法定代表人及负责人信息、联系人手机号和邮箱、主体授权书。"),
        ("应用材料", "App名称、简介、图标、截图、版本号、iOS Bundle ID、Android 包名、签名信息、功能说明、审核账号。"),
        ("网络材料", "域名证书、域名实名认证、云服务合同、服务器/IP信息、ICP备案信息、生产 API 和隐私政策网址。"),
        ("教育材料", "教育应用提供者信息、服务场景、学校合作或使用说明、教育主管部门要求的审核或备案材料。"),
        ("安全材料", "等保定级报告、备案证明、整改报告、测评或检测报告、应急预案、权限矩阵、日志和备份方案。"),
        ("隐私材料", "隐私政策、用户协议、儿童隐私规则、监护人同意记录、PIA、SDK清单、第三方处理协议、删除和注销流程。"),
        ("专项材料", "医疗器械分类意见、体育培训许可/登记、ICP许可证、EDI许可证、网络出版或视听许可（仅在触发时提供）。"),
        ("知识产权", "软件著作权登记证书、商标申请/注册材料、字体/图片/视频/课程内容授权、第三方 SDK 许可证明。"),
    ], [1800, 7560], font_size=9.7)

    add_heading(doc, "九、上线后持续管理", 1)
    for item in [
        "每次新增采集字段、第三方 SDK、支付、算法、公开内容或学校数据用途前，进行变更评估并更新隐私政策和备案信息。",
        "建立个人信息安全事件响应机制，明确发现、隔离、调查、通知、补救和监管报告时限。",
        "每年完成未成年人个人信息合规审计；按系统等级和主管部门要求开展等保复测或安全检查。",
        "对教师、学校管理员、客服和技术运维人员进行隐私与数据安全培训，保留培训和权限审计记录。",
        "持续核验课程、图片、视频、字体、AI模型和第三方 SDK 的版权、授权和许可范围。",
        "对学生风险标签、心理结果和困难地区标签实行最小可见、最短保存和人工复核，避免自动化结果直接影响评优、收费或教育机会。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十、最终结论", 1)
    add_p(doc, "本 App 可以上线，但不应按普通运动 App 直接发布。建议先完成“学校封闭试点版”的主体、域名、ICP/APP备案、教育 App 备案、等保、公安联网备案准备、未成年人隐私体系和非医疗定位整改，再逐项扩展商业化功能。", after=8)
    add_note(doc, "建议的上线门槛", "在取得基础备案、完成安全和隐私整改、可验证监护人授权、可用账号注销/删除、学校数据协议签署、生产后台可审计，并完成应用商店合规测试之前，不建议面向真实学生公开运营。", fill="EEF7F2", color=GREEN)

    doc.add_page_break()
    add_heading(doc, "附录：主要法规与官方入口", 1)
    sources = [
        ("[1] 工业和信息化部：移动互联网应用程序备案工作的通知", "https://www.miit.gov.cn/zwgk/zcwj/wjfb/tz/art/2023/art_920db564162e4312916a01bed6540ad8.html"),
        ("[2] 工业和信息化部：App备案政策解读（20个工作日）", "https://www.miit.gov.cn/jgsj/xgj/hlwgl/art/2023/art_564bf0759d7e41d5b4aa8ce4996b9e84.html"),
        ("[3] 教育部：教育移动互联网应用程序备案管理办法", "https://www.moe.gov.cn/srcsite/A16/s3342/201911/t20191122_409333.html?from=timeline&isappinstalled=0"),
        ("[4] 公安部：计算机信息网络国际联网安全保护管理办法", "https://www.miit.gov.cn/jgsj/xxjsfzs/xxgk/art/2020/art_25af2145083b40e98cc4993438703d78.html"),
        ("[5] 国家网信办：个人信息保护政策法规问答（2026年4月）", "https://www.cac.gov.cn/2026-04/29/c_1779200509387274.htm"),
        ("[6] 国家网信办：个人信息保护合规审计管理办法", "https://www.cac.gov.cn/2025-02/14/c_1741233507681519.htm"),
        ("[7] 国家市场监督管理总局：未成年人网络保护条例", "https://www.samr.gov.cn/wljys/gzzd/art/2024/art_43b2ac59aefd427aabe333df54d68f0a.html"),
        ("[8] 国家网信办：移动互联网应用程序信息服务管理规定", "https://www.cac.gov.cn/2022-06/14/c_1656821626455324.htm"),
        ("[9] 国家网信办：互联网信息服务算法推荐管理规定", "https://www.cac.gov.cn/2022-01/04/c_1642894606364259.htm"),
        ("[10] 国家市场监督管理总局：医疗器械监督管理条例", "https://www.samr.gov.cn/zw/zfxxgk/fdzdgknr/fgs/art/2023/art_70607fc4160041a383e68ff6bfb2826f.html"),
        ("[11] 国家体育总局：体育类课外培训机构监管答复", "https://www.sport.gov.cn/n323/n10459/c28669965/content.html"),
        ("[12] 国家新闻出版署：网络出版服务管理规定", "https://www.nppa.gov.cn/xxfb/zcfg/bmgz/201602/t20160206_4403.html"),
    ]
    for title, url in sources:
        p = doc.add_paragraph()
        set_para(p, 0, 2, 1.05)
        r = p.add_run(title + "\n")
        set_run_font(r, size=8.5, color=NAVY, bold=True)
        r2 = p.add_run(url)
        set_run_font(r2, name="Aptos", size=7.5, color=BLUE)

    add_p(doc, "说明：本报告为项目上线合规分析和办理清单，不替代律师出具的专项法律意见，也不替代通信、教育、公安、网信、市场监管、体育或药监部门的正式认定。", before=12, after=0, size=9.5, color=MUTED, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
